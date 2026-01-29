"""
HTML 轉換核心模組
負責將 Word 文件轉換為 HTML,包含自動目錄生成、平滑捲動等功能
"""

from docx import Document
from html import escape
import re
from .word_processor import (
    remove_content_controls,
    paragraph_to_html_with_links,
    table_to_html,
    is_pure_url,
    extract_embed_url,
    convert_url_to_iframe
)
from templates.footers import CATEGORY_TO_FOOTER_HTML


def apply_auto_toc_and_smooth(html_list):
    """
    自動生成目錄並加入平滑捲動效果
    目錄只插在第一個 H2 前
    
    Args:
        html_list: HTML 區塊列表
        
    Returns:
        list: 處理後的 HTML 區塊列表
    """
    updated = []
    toc = []
    first_h2_index = None
    
    # 加入平滑捲動 CSS
    smooth_css = """<style>html { scroll-behavior: smooth; } .html-container h2, .html-container h3 { scroll-margin-top: 130px; }</style>"""
    updated.append(smooth_css)
    
    for block in html_list:
        b = (block or "").strip()
        plain_text = re.sub(r"<[^>]*>", "", b) if b else ""
        
        if b.startswith("<h1"):
            anchor = f"toc-h1-{len(toc)}"
            toc.append((1, plain_text, anchor))
            block = block.replace("<h1>", f"<h1 id='{anchor}'>", 1)
        
        elif b.startswith("<h2"):
            anchor = f"toc-h2-{len(toc)}"
            toc.append((2, plain_text, anchor))
            block = block.replace("<h2", f"<h2 id='{anchor}'", 1)
            if first_h2_index is None:
                first_h2_index = len(updated)
        
        elif b.startswith("<h3"):
            anchor = f"toc-h3-{len(toc)}"
            toc.append((3, plain_text, anchor))
            block = block.replace("<h3", f"<h3 id='{anchor}'", 1)
        
        updated.append(block)
    
    if first_h2_index is None or not toc:
        return updated
    
    def is_blank_para(x: str):
        s = (x or "").strip()
        return s in ("<p>&nbsp;</p>", "<p>&nbsp;</p><p>&nbsp;</p>")
    
    # 移除目錄前的空白段落
    while first_h2_index - 1 >= 0 and is_blank_para(updated[first_h2_index - 1]):
        updated.pop(first_h2_index - 1)
        first_h2_index -= 1
    
    # 生成目錄 HTML
    toc_html = [
        "<div style='margin-top:28px; margin-bottom:12px; padding:12px 0 12px 16px; border-left:4px solid #4f8ef7;'>",
        "<div style='font-size:20px; font-weight:700; margin-bottom:10px; color:#000000;'>文章目錄</div>",
        "<ul style='list-style:none; margin-left: 28px; padding-left: 0; line-height:1.8; font-size:17px; color:#4f8ef7;'>"
    ]
    
    for level, text, anchor in toc:
        safe_text = escape(text) if text else ""
        if level == 2:
            bullet = "•"
            indent_px = 0
        elif level == 3:
            bullet = "◦"
            indent_px = 18
        else:
            bullet = "•"
            indent_px = 0
        
        toc_html.append(
            f"<li style='margin:6px 0; padding-left:{indent_px}px; text-indent:-12px;'>"
            f"<span style='display:inline-block; width:12px; opacity:0.7;'>{bullet}</span>"
            f"<a href='#{anchor}' style='color:#4f8ef7; text-decoration:none;'>{safe_text}</a>"
            f"</li>"
        )
    
    toc_html.append("</ul></div>")
    
    # 插入目錄
    updated = updated[:first_h2_index] + toc_html + updated[first_h2_index:]
    insert_after_toc = first_h2_index + len(toc_html)
    updated = updated[:insert_after_toc] + ["<p>&nbsp;</p><p>&nbsp;</p>"] + updated[insert_after_toc:]
    
    return updated


def docx_to_html_with_links(doc_path: str, category_choice: str):
    """
    將 Word 文件轉換為 HTML
    
    Args:
        doc_path: Word 文件路徑
        category_choice: 文章分類 (決定使用哪個 Footer)
        
    Returns:
        tuple: (html_content, h1_title)
            - html_content: 完整的 HTML 內容
            - h1_title: 文章主標題 (H1)
    """
    doc = Document(doc_path)
    remove_content_controls(doc)
    html_output = []
    h1_text = None
    last_was_blank = False
    elements = list(doc.element.body)
    total = len(elements)
    
    for idx, element in enumerate(elements):
        tag = element.tag.split('}')[-1]
        next_tag = elements[idx + 1].tag.split('}')[-1] if idx < total - 1 else None
        next_style = ""
        if next_tag == "p":
            for np in doc.paragraphs:
                if np._element == elements[idx + 1]:
                    next_style = np.style.name.lower() if np.style and np.style.name else ''
                    break
        
        if tag == "p":
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if not para:
                continue
            
            text = (para.text or "").strip()
            style = para.style.name.lower() if para.style and para.style.name else ''
            
            if not text:
                continue
            
            # 處理「整段只有 URL」→ 嘗試轉成 iframe
            if is_pure_url(text):
                url = extract_embed_url(text) or text
                iframe_block = convert_url_to_iframe(url)
                if iframe_block:
                    # 確保 iframe 上方只有 1 個空行
                    if not last_was_blank:
                        html_output.append("<p>&nbsp;</p>")
                    
                    html_output.append(iframe_block)
                    
                    # iframe 下方固定 1 行空行
                    html_output.append("<p>&nbsp;</p>")
                    
                    last_was_blank = True
                    continue
            
            # 一般段落轉換
            content_html = paragraph_to_html_with_links(para)
            
            # 移除 Word 自帶 TOC 的編號
            if style.startswith("toc"):
                content_html = re.sub(r"^\s*[\d\.\-\(\)、．]+\s*", "", content_html)
            
            # H1 當作主標題,只抓文字不輸出
            if 'heading 1' in style:
                if not h1_text:
                    h1_text = text
                continue
            
            # H2
            if 'heading 2' in style:
                html_output.append('<p>&nbsp;</p>' if last_was_blank else '<p>&nbsp;</p>' * 2)
                html_output.append(
                    f'<h2 style="padding-top:150px; margin-top:-150px;">'
                    f'<span style="color:#0066CC;"><span style="font-size:20px;"><strong>{content_html}</strong></span></span>'
                    f'</h2>'
                )
                html_output.append('<p>&nbsp;</p>')
                last_was_blank = True
                continue
            
            # H3
            if 'heading 3' in style:
                if not last_was_blank:
                    html_output.append('<p>&nbsp;</p>')
                html_output.append(
                    f'<h3 style="padding-top:150px; margin-top:-150px;">'
                    f'<span style="color:#000000;"><span style="font-size:18px;"><strong>{content_html}</strong></span></span>'
                    f'</h3>'
                )
                html_output.append('<p>&nbsp;</p>')
                last_was_blank = True
                continue
            
            # 一般內文
            html_output.append(f'<p>{content_html}</p>')
            if not (('heading 2' in next_style) or ('heading 3' in next_style) or (next_tag == "tbl")):
                html_output.append('<p>&nbsp;</p>')
                last_was_blank = True
            else:
                last_was_blank = False
        
        elif tag == "tbl":
            for tbl in doc.tables:
                if tbl._element == element:
                    if not last_was_blank:
                        html_output.append('<p>&nbsp;</p>')
                    html_output.append(table_to_html(tbl))
                    html_output.append('<p>&nbsp;</p>')
                    last_was_blank = True
                    break
    
    # 加入 Footer
    footer_html = CATEGORY_TO_FOOTER_HTML.get(category_choice, "")
    html_output.append('<p>&nbsp;</p><p>&nbsp;</p>' + footer_html + '<p>&nbsp;</p>')
    
    # 應用自動目錄和平滑捲動
    html_output = apply_auto_toc_and_smooth(html_output)
    
    result = "\n".join(html_output)
    
    return result, h1_text or "（文件中無 H1 標題）"
